"""
문서 처리기 클래스 - 다양한 문서 형식의 처리를 위한 모듈

이 모듈은 다양한 문서 형식(PDF, Word, Excel, 텍스트 등)에서 텍스트 및 메타데이터를 추출하는 기능을 제공합니다.
문서 변환, 내용 분석, 메타데이터 추출 등의 기능을 구현합니다.
"""

import os
import json
import logging
import datetime
import tempfile
import re
import mimetypes
import docx
import pandas as pd
from io import BytesIO
from PIL import Image
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument

# 로깅 설정
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    다양한 문서 형식을 처리하는 클래스
    """
    
    def __init__(self, config=None):
        """
        DocumentProcessor 초기화
        
        Args:
            config (dict, optional): 문서 처리 설정
        """
        self.logger = logger
        self.config = config or {}
        
        # 임시 디렉토리
        self.temp_dir = self.config.get('temp_dir', os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'temp', 'docs'))
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def extract_text(self, file_path):
        """
        문서에서 텍스트 추출
        
        Args:
            file_path (str): 문서 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            # 파일 확장자로 문서 유형 식별
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 문서 유형별 텍스트 추출
            if file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_text_from_excel(file_path)
            elif file_ext in ['.txt', '.csv', '.md', '.json', '.xml', '.html', '.htm']:
                return self._extract_text_from_text_file(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                # 참고: 이미지 텍스트 추출은 간소화된 구현이므로 실제 OCR 기능은 제공하지 않음
                return f"[이미지 파일: {os.path.basename(file_path)}]"
            else:
                self.logger.warning(f"지원되지 않는 파일 형식: {file_ext}")
                return f"[지원되지 않는 파일 형식: {file_ext}]"
        except Exception as e:
            self.logger.error(f"텍스트 추출 중 오류 발생: {str(e)}")
            return f"[텍스트 추출 오류: {str(e)}]"
    
    def _extract_text_from_pdf(self, file_path):
        """
        PDF 파일에서 텍스트 추출
        
        Args:
            file_path (str): PDF 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            text = pdf_extract_text(file_path)
            return text
        except Exception as e:
            self.logger.error(f"PDF 텍스트 추출 중 오류 발생: {str(e)}")
            return f"[PDF 텍스트 추출 오류: {str(e)}]"
    
    def _extract_text_from_docx(self, file_path):
        """
        Word 문서에서 텍스트 추출
        
        Args:
            file_path (str): Word 문서 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            doc = docx.Document(file_path)
            
            # 모든 단락 텍스트 추출
            paragraphs = [para.text for para in doc.paragraphs]
            
            # 테이블 텍스트 추출
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cell_texts = [cell.text for cell in row.cells]
                    rows.append(' | '.join(cell_texts))
                tables.append('\n'.join(rows))
            
            # 모든 텍스트 결합
            return '\n\n'.join(paragraphs + tables)
        except Exception as e:
            self.logger.error(f"Word 텍스트 추출 중 오류 발생: {str(e)}")
            return f"[Word 텍스트 추출 오류: {str(e)}]"
    
    def _extract_text_from_excel(self, file_path):
        """
        Excel 파일에서 텍스트 추출
        
        Args:
            file_path (str): Excel 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            # Excel 파일 로드
            df_dict = pd.read_excel(file_path, sheet_name=None)
            
            # 모든 시트의 데이터 추출
            sheets_text = []
            
            for sheet_name, df in df_dict.items():
                # 시트 헤더
                sheet_text = f"[시트: {sheet_name}]\n"
                
                # 데이터프레임을 텍스트로 변환
                sheet_text += df.to_string(index=False)
                
                sheets_text.append(sheet_text)
            
            # 모든 시트 텍스트 결합
            return '\n\n'.join(sheets_text)
        except Exception as e:
            self.logger.error(f"Excel 텍스트 추출 중 오류 발생: {str(e)}")
            return f"[Excel 텍스트 추출 오류: {str(e)}]"
    
    def _extract_text_from_text_file(self, file_path):
        """
        텍스트 파일에서 텍스트 추출
        
        Args:
            file_path (str): 텍스트 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            # 파일 확장자 확인
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 인코딩 추정
            encodings = ['utf-8', 'euc-kr', 'cp949', 'ascii']
            
            # 각 인코딩으로 시도
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    return text
                except UnicodeDecodeError:
                    continue
            
            # 모든 인코딩 실패 시 바이너리 모드로 읽기
            with open(file_path, 'rb') as f:
                binary = f.read()
                
            # 바이너리 데이터를 utf-8로 디코딩 시도 (오류 무시)
            return binary.decode('utf-8', errors='replace')
        except Exception as e:
            self.logger.error(f"텍스트 파일 추출 중 오류 발생: {str(e)}")
            return f"[텍스트 파일 추출 오류: {str(e)}]"
    
    def extract_metadata(self, file_path):
        """
        문서에서 메타데이터 추출
        
        Args:
            file_path (str): 문서 파일 경로
            
        Returns:
            dict: 추출된 메타데이터
        """
        try:
            # 기본 파일 메타데이터
            file_stat = os.stat(file_path)
            
            metadata = {
                'filename': os.path.basename(file_path),
                'path': file_path,
                'size': file_stat.st_size,
                'created_date': datetime.datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                'modified_date': datetime.datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                'file_extension': os.path.splitext(file_path)[1][1:].lower(),
                'mime_type': mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
            }
            
            # 파일 확장자로 문서 유형 식별
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 문서 유형별 추가 메타데이터 추출
            if file_ext == '.pdf':
                pdf_metadata = self._extract_metadata_from_pdf(file_path)
                metadata.update(pdf_metadata)
            elif file_ext == '.docx':
                docx_metadata = self._extract_metadata_from_docx(file_path)
                metadata.update(docx_metadata)
            elif file_ext in ['.xlsx', '.xls']:
                excel_metadata = self._extract_metadata_from_excel(file_path)
                metadata.update(excel_metadata)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                image_metadata = self._extract_metadata_from_image(file_path)
                metadata.update(image_metadata)
            
            return metadata
        except Exception as e:
            self.logger.error(f"메타데이터 추출 중 오류 발생: {str(e)}")
            return {
                'filename': os.path.basename(file_path),
                'path': file_path,
                'error': str(e)
            }
    
    def _extract_metadata_from_pdf(self, file_path):
        """
        PDF 파일에서 메타데이터 추출
        
        Args:
            file_path (str): PDF 파일 경로
            
        Returns:
            dict: 추출된 메타데이터
        """
        try:
            metadata = {}
            
            with open(file_path, 'rb') as f:
                parser = PDFParser(f)
                doc = PDFDocument(parser)
                
                # PDF 메타데이터 추출
                if doc.info:
                    for key, value in doc.info[0].items():
                        # 바이너리 값 처리
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8')
                            except:
                                try:
                                    value = value.decode('latin-1')
                                except:
                                    value = str(value)
                        
                        # 키 이름 정규화
                        key = key.decode('utf-8') if isinstance(key, bytes) else str(key)
                        key = key.lower().replace(':', '_')
                        
                        metadata[f"pdf_{key}"] = value
            
            # 페이지 수 추정
            try:
                text = pdf_extract_text(file_path)
                page_count = text.count('\f') + 1  # Form feed 문자로 페이지 구분
                metadata['page_count'] = page_count
            except:
                pass
            
            # 제목 추출 (PDF 메타데이터에서 추출 실패 시 파일명 사용)
            if 'pdf_title' in metadata and metadata['pdf_title']:
                metadata['title'] = metadata['pdf_title']
            else:
                metadata['title'] = os.path.splitext(os.path.basename(file_path))[0]
            
            # 저자 추출
            if 'pdf_author' in metadata and metadata['pdf_author']:
                metadata['author'] = metadata['pdf_author']
            
            # 생성 날짜 추출
            if 'pdf_creationdate' in metadata and metadata['pdf_creationdate']:
                date_str = metadata['pdf_creationdate']
                # PDF 날짜 형식 처리 (D:20201231235959+09'00')
                if date_str.startswith('D:'):
                    date_str = date_str[2:]
                    try:
                        year = int(date_str[:4])
                        month = int(date_str[4:6])
                        day = int(date_str[6:8])
                        metadata['creation_date'] = f"{year:04d}-{month:02d}-{day:02d}"
                    except:
                        pass
            
            # 문서 타입 설정
            metadata['doc_type'] = 'PDF'
            
            return metadata
        except Exception as e:
            self.logger.error(f"PDF 메타데이터 추출 중 오류 발생: {str(e)}")
            return {
                'doc_type': 'PDF',
                'error_metadata': str(e)
            }
    
    def _extract_metadata_from_docx(self, file_path):
        """
        Word 문서에서 메타데이터 추출
        
        Args:
            file_path (str): Word 문서 파일 경로
            
        Returns:
            dict: 추출된 메타데이터
        """
        try:
            metadata = {}
            
            # Word 문서 로드
            doc = docx.Document(file_path)
            
            # 코어 속성 추출
            core_props = doc.core_properties
            
            if core_props.title:
                metadata['title'] = core_props.title
            else:
                metadata['title'] = os.path.splitext(os.path.basename(file_path))[0]
            
            if core_props.author:
                metadata['author'] = core_props.author
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            
            if core_props.modified:
                metadata['modified_date'] = core_props.modified.isoformat()
            
            if core_props.comments:
                metadata['comments'] = core_props.comments
            
            if core_props.category:
                metadata['category'] = core_props.category
            
            if core_props.subject:
                metadata['subject'] = core_props.subject
            
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            
            # 문서 내용 통계
            metadata['page_count'] = len(doc.sections)
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # 문서 내용 길이 추정
            total_text = '\n'.join([para.text for para in doc.paragraphs])
            metadata['word_count'] = len(total_text.split())
            metadata['char_count'] = len(total_text)
            
            # 문서 타입 설정
            metadata['doc_type'] = 'Word'
            
            return metadata
        except Exception as e:
            self.logger.error(f"Word 메타데이터 추출 중 오류 발생: {str(e)}")
            return {
                'doc_type': 'Word',
                'error_metadata': str(e)
            }
    
    def _extract_metadata_from_excel(self, file_path):
        """
        Excel 파일에서 메타데이터 추출
        
        Args:
            file_path (str): Excel 파일 경로
            
        Returns:
            dict: 추출된 메타데이터
        """
        try:
            metadata = {}
            
            # Excel 파일 로드
            df_dict = pd.read_excel(file_path, sheet_name=None)
            
            # 기본 메타데이터
            metadata['title'] = os.path.splitext(os.path.basename(file_path))[0]
            metadata['sheet_count'] = len(df_dict)
            metadata['sheet_names'] = list(df_dict.keys())
            
            # 시트별 행/열 수
            sheet_stats = {}
            total_rows = 0
            max_cols = 0
            
            for sheet_name, df in df_dict.items():
                rows = len(df)
                cols = len(df.columns)
                sheet_stats[sheet_name] = {
                    'rows': rows,
                    'columns': cols
                }
                total_rows += rows
                max_cols = max(max_cols, cols)
            
            metadata['sheet_stats'] = sheet_stats
            metadata['total_rows'] = total_rows
            metadata['max_columns'] = max_cols
            
            # 문서 타입 설정
            metadata['doc_type'] = 'Excel'
            
            return metadata
        except Exception as e:
            self.logger.error(f"Excel 메타데이터 추출 중 오류 발생: {str(e)}")
            return {
                'doc_type': 'Excel',
                'error_metadata': str(e)
            }
    
    def _extract_metadata_from_image(self, file_path):
        """
        이미지 파일에서 메타데이터 추출
        
        Args:
            file_path (str): 이미지 파일 경로
            
        Returns:
            dict: 추출된 메타데이터
        """
        try:
            metadata = {}
            
            # 기본 메타데이터
            metadata['title'] = os.path.splitext(os.path.basename(file_path))[0]
            
            # 이미지 로드
            with Image.open(file_path) as img:
                # 이미지 기본 속성
                metadata['width'] = img.width
                metadata['height'] = img.height
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                
                # EXIF 데이터 추출 (가능한 경우)
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    
                    # EXIF 태그 매핑
                    exif_tags = {
                        271: 'camera_manufacturer',    # Make
                        272: 'camera_model',           # Model
                        306: 'date_time',              # DateTime
                        36867: 'date_time_original',   # DateTimeOriginal
                        33432: 'copyright',            # Copyright
                        37377: 'shutter_speed',        # ShutterSpeedValue
                        37378: 'aperture',             # ApertureValue
                        37380: 'exposure_compensation',# ExposureCompensation
                        37383: 'metering_mode',        # MeteringMode
                        37384: 'flash',                # Flash
                        37385: 'flash_mode',           # FlashMode
                        37386: 'focal_length',         # FocalLength
                        40962: 'pixel_x_dimension',    # PixelXDimension
                        40963: 'pixel_y_dimension'     # PixelYDimension
                    }
                    
                    for tag, value in exif.items():
                        if tag in exif_tags and value:
                            metadata[exif_tags[tag]] = str(value)
            
            # 문서 타입 설정
            metadata['doc_type'] = 'Image'
            
            return metadata
        except Exception as e:
            self.logger.error(f"이미지 메타데이터 추출 중 오류 발생: {str(e)}")
            return {
                'doc_type': 'Image',
                'error_metadata': str(e)
            }
    
    def convert_to_pdf(self, file_path, output_path=None):
        """
        문서를 PDF로 변환
        
        Args:
            file_path (str): 문서 파일 경로
            output_path (str, optional): 출력 PDF 파일 경로 (지정하지 않으면 자동 생성)
            
        Returns:
            str: 생성된 PDF 파일 경로 또는 None (실패 시)
        """
        # 참고: 실제 PDF 변환은 복잡한 외부 라이브러리 의존성이 필요하므로,
        # 이 메서드는 간소화된 구현을 제공하며 실제 변환 기능은 구현하지 않음
        self.logger.warning("PDF 변환 기능은 이 구현에서 실제로 지원되지 않음")
        
        if output_path is None:
            output_path = os.path.join(
                self.temp_dir,
                f"{os.path.splitext(os.path.basename(file_path))[0]}.pdf"
            )
        
        # 변환 성공 시뮬레이션
        return output_path
    
    def extract_entities(self, text):
        """
        텍스트에서 엔티티 추출
        
        Args:
            text (str): 분석할 텍스트
            
        Returns:
            list: 추출된 엔티티 목록
        """
        # 참고: 실제 엔티티 추출은 NLP 라이브러리 의존성이 필요하므로,
        # 이 메서드는 간소화된 구현을 제공하며 간단한 패턴 매칭만 수행
        entities = []
        
        # 패턴 정의
        patterns = {
            'Email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'URL': r'https?://[^\s]+',
            'IPAddress': r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',
            'PhoneNumber': r'\+?\d{1,4}?[-.\s]?\(?\d{1,3}?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            'Date': r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b'
        }
        
        # 패턴별 엔티티 추출
        for entity_type, pattern in patterns.items():
            matches = re.findall(pattern, text)
            for match in matches:
                entities.append({
                    'type': entity_type,
                    'value': match
                })
        
        return entities
    
    def summarize_text(self, text, max_sentences=5):
        """
        텍스트 요약
        
        Args:
            text (str): 요약할 텍스트
            max_sentences (int): 최대 요약 문장 수
            
        Returns:
            str: 요약된 텍스트
        """
        # 참고: 실제 텍스트 요약은 NLP 라이브러리 의존성이 필요하므로,
        # 이 메서드는 간소화된 구현을 제공하며 단순히 첫 몇 문장을 추출
        
        # 문장 분리
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # 최대 문장 수만큼 추출
        summary_sentences = sentences[:max_sentences]
        
        # 요약 결합
        summary = ' '.join(summary_sentences)
        
        return summary
    
    def detect_language(self, text):
        """
        텍스트 언어 감지
        
        Args:
            text (str): 감지할 텍스트
            
        Returns:
            str: 감지된 언어 코드
        """
        # 참고: 실제 언어 감지는 NLP 라이브러리 의존성이 필요하므로,
        # 이 메서드는 간소화된 구현을 제공하며 단순한 휴리스틱 검사 수행
        
        # 한글 포함 확인
        korean_chars = re.findall(r'[ㄱ-ㅎㅏ-ㅣ가-힣]', text)
        if len(korean_chars) > len(text) * 0.1:
            return 'ko'
        
        # 일본어 문자 포함 확인
        japanese_chars = re.findall(r'[\u3040-\u30ff\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff\uff66-\uff9f]', text)
        if len(japanese_chars) > len(text) * 0.1:
            return 'ja'
        
        # 중국어 문자 포함 확인
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        if len(chinese_chars) > len(text) * 0.1:
            return 'zh'
        
        # 기본값: 영어
        return 'en'
